from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Flick Series_MVP_기능기획서.docx"
FONT_NAME = "AppleGothic"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def format_table(table, widths):
    table.style = "Table Grid"
    table.autofit = False
    set_table_width(table, sum(widths))
    set_cell_margins(table)
    for row_index, row in enumerate(table.rows):
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(4)
                for run in paragraph.runs:
                    run.font.name = FONT_NAME
                    rfonts = run._element.rPr.rFonts
                    rfonts.set(qn("w:ascii"), FONT_NAME)
                    rfonts.set(qn("w:hAnsi"), FONT_NAME)
                    rfonts.set(qn("w:eastAsia"), FONT_NAME)
                    rfonts.set(qn("w:cs"), FONT_NAME)
                    run.font.size = Pt(10)
            if row_index == 0:
                set_cell_shading(cell, "F8F9FA")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def set_font(run, size=None, bold=None, color=None):
    run.font.name = FONT_NAME
    rfonts = run._element.rPr.rFonts
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", style=None, after=8):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if text:
        r = p.add_run(text)
        set_font(r, 11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r, 11)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r, 11)
    return p


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
    else:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        set_font(run, {1: 20, 2: 16, 3: 14}.get(level, 12), bold=False, color="000000" if level < 3 else "434343")
    return p


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "항목"
    table.rows[0].cells[1].text = "내용"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    format_table(table, [2160, 7200])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_matrix_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = value
    format_table(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal_rfonts = normal._element.rPr.rFonts
    normal_rfonts.set(qn("w:ascii"), FONT_NAME)
    normal_rfonts.set(qn("w:hAnsi"), FONT_NAME)
    normal_rfonts.set(qn("w:eastAsia"), FONT_NAME)
    normal_rfonts.set(qn("w:cs"), FONT_NAME)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Flick Series MVP 기능 기획서")
    set_font(run, 26, bold=False, color="000000")

    subtitle = add_para(doc, "노트북 힌지 제스처 기반 macOS 유틸리티 시리즈", after=8)
    subtitle.runs[0].font.color.rgb = RGBColor(85, 85, 85)

    meta = add_para(doc, "문서 유형: MVP 확장 기능 기획서  |  버전: v1.2  |  작성일: 2026.07.17  |  상태: Flick Privacy / Cancel 확정본", after=14)
    meta.runs[0].font.color.rgb = RGBColor(85, 85, 85)

    add_heading(doc, "1. 문서 목적", 1)
    add_para(
        doc,
        "본 문서는 Flick Series MVP 확장 범위에서 신규 기능인 Flick Privacy와 복구 기능인 Flick Privacy Cancel의 제품 의도, 사용자 가치, 세부 기능, 설정 범위, 구현 제약, 검수 기준을 정리한 기능 기획서이다. 개발, 디자인, QA, 데모 준비 담당자가 동일한 기준으로 기능을 이해하고 납품 가능 여부를 판단할 수 있도록 작성한다.",
    )

    add_heading(doc, "2. 제품 개요", 1)
    add_para(
        doc,
        "Flick Series는 노트북 힌지, 즉 Lid의 움직임 패턴을 입력 장치로 활용하는 macOS 유틸리티 시리즈이다. 사용자는 마우스 조작이나 키보드 단축키 없이 노트북 화면을 특정 방식으로 움직여 창 정리, 집중 모드, 프라이버시 보호 같은 작업 환경 제어 기능을 실행한다.",
    )
    add_para(doc, "핵심 제품 철학은 다음과 같다.", after=4)
    add_bullet(doc, "Flick Series는 일반적인 창 관리 앱이 아니다.")
    add_bullet(doc, "핵심 가설은 노트북 힌지를 자연스러운 입력 장치로 사용할 수 있는지 검증하는 것이다.")
    add_bullet(doc, "MVP 확장 기능은 “노트북을 닫는 동작”만으로 공공장소, 회의실, 이동 전 상황에서 화면 노출과 소리 노출을 즉시 줄일 수 있어야 한다.")

    add_heading(doc, "3. 명칭 및 입력 규칙", 1)
    add_para(doc, "입력 제스처와 앱 기능명은 서로 다른 명명 규칙을 따른다. 이 규칙은 UI 문구, 로그, 문서, 데모 스크립트에서 일관되게 적용한다.")
    add_matrix_table(
        doc,
        ["구분", "규칙", "예시"],
        [
            ["제스처 입력", "Flick이라는 단어를 사용하지 않는다. 사용 가능한 제스처는 Close와 Open 두 개뿐이다.", "Close, Open"],
            ["앱 기능", "기존처럼 Flick 접두어를 유지한다.", "Flick Arrange, Flick Privacy, Flick Privacy Cancel"],
        ],
        [1800, 5400, 2160],
    )

    add_heading(doc, "4. 제스처 정의", 1)
    add_matrix_table(
        doc,
        ["제스처", "입력 패턴", "설명", "MVP 사용 여부"],
        [
            ["Close", "110° → 90° → 97°", "화면을 빠르게 닫았다가 약간 복원하는 동작", "사용"],
            ["Open", "100° → 120° → 113°", "화면을 빠르게 열었다가 약간 복원하는 동작", "Flick Privacy Cancel 기본 할당"],
        ],
        [1200, 2200, 4160, 1800],
    )
    add_para(doc, "현재 MVP 확장 범위에서는 Close 제스처에 Flick Privacy를 할당하고, Close에 Flick Privacy가 할당된 경우 Open 제스처에는 기본적으로 Flick Privacy Cancel을 자동 할당한다. 단, 사용자는 Open에 할당된 기능을 자유롭게 다른 기능으로 수정할 수 있어야 한다.")

    add_heading(doc, "5. MVP 기능 구성", 1)
    add_matrix_table(
        doc,
        ["우선순위", "기능명", "트리거", "상태", "핵심 가치"],
        [
            ["P0", "Flick Arrange", "Close", "완료", "복잡해진 작업 공간을 즉시 정리한다."],
            ["P1", "Flick Privacy", "Close", "신규 구현 확정", "밝기, 소리, 민감 앱, 업무 모드, 미디어, 창 숨김을 한 번에 적용한다."],
            ["P1", "Flick Privacy Cancel", "Open", "신규 구현 확정", "Flick Privacy 실행 이전 환경을 가능한 한 정확히 복구한다."],
        ],
        [1100, 2100, 1300, 1500, 3360],
    )

    add_heading(doc, "6. 기능 요구사항", 1)
    add_heading(doc, "6.1 Flick Arrange", 2)
    add_kv_table(
        doc,
        [
            ("트리거", "Close"),
            ("현재 상태", "구현 완료"),
            ("기능 설명", "현재 열려 있는 창들을 자동 정렬한다."),
            ("사용자 목적", "복잡해진 작업 공간을 즉시 정리한다."),
            ("MVP 기준", "기존 구현을 유지하며 신규 기능과 명칭 규칙만 충돌 없이 정리한다."),
        ],
    )

    add_heading(doc, "6.2 Flick Privacy", 2)
    add_kv_table(
        doc,
        [
            ("트리거", "Close"),
            ("목적", "노트북 환경에서 화면 노출, 소리 노출, 민감 앱 노출을 즉시 줄여 프라이버시 작업 상태로 전환한다."),
            ("핵심 동작", "사용자가 켠 세부 기능을 순차 실행해 밝기, 볼륨, 민감 앱, 업무 모드, 미디어 재생, 창 표시 상태를 제어한다."),
            ("사용 맥락", "카페, 회의실, 도서관, 이동 직전처럼 노트북 화면과 소리가 주변에 노출될 수 있는 상황"),
            ("MVP 기준", "Flick Privacy를 신규 확정 기능으로 구현하며, Flick Focus는 Privacy 내부 세부 기능으로 재사용한다. Privacy 실행 전 상태는 Flick Privacy Cancel이 복구할 수 있도록 스냅샷으로 저장한다."),
        ],
    )
    add_heading(doc, "세부 기능 설정", 3)
    add_matrix_table(
        doc,
        ["번호", "세부 기능", "기본값", "설정 방식"],
        [
            ["1", "화면 밝기를 전체 밝기 기준 30%로 낮추기", "ON / 30%", "토글 및 0~100% 슬라이더"],
            ["2", "시스템 볼륨을 0%로 만들기", "ON", "토글"],
            ["3", "메신저 및 메일 앱 숨기기", "ON", "기본 앱 목록 + 사용자 추가/삭제 목록"],
            ["4", "업무 모드로 전환하기", "ON", "토글"],
            ["5", "현재 시스템 미디어 재생 일시정지", "ON", "토글"],
            ["6", "현재 창 외 나머지 숨김", "ON", "토글 / Flick Focus 로직 재사용"],
        ],
        [800, 4200, 1560, 2800],
    )
    add_heading(doc, "기본 숨김 대상 앱", 3)
    add_para(doc, "메신저/메일 앱 숨김은 macOS가 앱을 의미론적으로 자동 분류하는 방식이 아니라, 번들 ID 또는 실행 앱 이름 목록 기반으로 구현한다. 사용자는 Privacy 설정에서 숨김 대상 앱을 추가하거나 삭제할 수 있어야 한다.")
    add_matrix_table(
        doc,
        ["분류", "기본 대상"],
        [
            ["메신저", "KakaoTalk, Messages/iMessage, Slack, Discord, Telegram, WhatsApp, Messenger"],
            ["메일", "Mail / Apple Mail / iMail"],
            ["사용자 지정", "사용자가 Privacy 설정에서 실행 중인 일반 앱 목록을 기준으로 추가 또는 삭제"],
        ],
        [1800, 7560],
    )
    add_heading(doc, "동작 원칙", 3)
    add_bullet(doc, "Flick Privacy 실행 시 켜져 있는 세부 기능만 순차 실행한다.")
    add_bullet(doc, "한 단계가 실패해도 전체 실행을 중단하지 않고 가능한 나머지 단계를 계속 실행한다.")
    add_bullet(doc, "실행 결과는 최근 활동 로그 또는 상태 메시지에 요약한다.")
    add_bullet(doc, "밝기, 업무 모드, 미디어 일시정지처럼 환경에 따라 실패할 수 있는 기능은 독립적으로 실패 처리한다.")
    add_heading(doc, "구현 제약", 3)
    add_bullet(doc, "밝기는 현재 밝기의 30%가 아니라 전체 밝기 스케일 기준 기본 30%로 설정한다.")
    add_bullet(doc, "업무 모드 전환은 macOS Focus / Do Not Disturb / Shortcuts / AppleScript 등 가능한 방식으로 검토하되, OS 권한과 버전에 따라 동작이 제한될 수 있다.")
    add_bullet(doc, "미디어 일시정지는 재생 중인 앱과 콘텐츠를 완벽히 식별하는 기능이 아니라, 시스템 미디어 Pause 명령을 보내는 방식으로 정의한다.")
    add_bullet(doc, "현재 창 외 나머지 숨김은 기존 Flick Focus 기능을 그대로 재사용한다.")
    add_heading(doc, "수용 기준", 3)
    add_number(doc, "Close 제스처가 인식되면 Flick Privacy가 실행된다.")
    add_number(doc, "각 세부 기능은 앱 UI에서 독립적으로 on/off 할 수 있다.")
    add_number(doc, "밝기 기능이 켜져 있으면 내장 디스플레이 밝기가 기본 30% 또는 사용자가 설정한 값으로 조정된다.")
    add_number(doc, "볼륨 기능이 켜져 있으면 시스템 볼륨이 0%가 된다.")
    add_number(doc, "메신저/메일 숨김 기능이 켜져 있으면 기본 대상 앱과 사용자 지정 앱이 화면에서 숨겨진다.")
    add_number(doc, "업무 모드 기능이 켜져 있으면 가능한 방식으로 업무 모드 전환을 시도하고, 실패 시 사용자에게 제한을 안내한다.")
    add_number(doc, "미디어 기능이 켜져 있으면 현재 시스템 미디어 재생에 Pause 명령을 보낸다.")
    add_number(doc, "창 숨김 기능이 켜져 있으면 활성 창을 유지하고 나머지 일반 앱 창은 숨긴다.")

    add_heading(doc, "6.3 Flick Privacy Cancel", 2)
    add_kv_table(
        doc,
        [
            ("트리거", "Open 기본 할당"),
            ("목적", "Flick Privacy 실행 이전의 작업 환경을 가능한 한 정확히 복구한다."),
            ("핵심 동작", "Flick Privacy 실행 직전에 저장한 밝기, 볼륨, 업무 모드, 숨김 앱, 활성 앱/창 상태를 기준으로 복구를 시도한다."),
            ("UI 정책", "Close에 Flick Privacy가 할당되면 Open에는 기본적으로 Flick Privacy Cancel을 자동 할당한다."),
            ("사용자 수정", "사용자는 Open에 할당된 Flick Privacy Cancel을 다른 기능으로 자유롭게 변경할 수 있다."),
        ],
    )
    add_heading(doc, "복구 대상", 3)
    add_matrix_table(
        doc,
        ["복구 항목", "복구 기준", "제약"],
        [
            ["밝기", "Privacy 실행 전 밝기 값으로 복구", "디스플레이 제어 가능 환경에서만 보장"],
            ["볼륨", "Privacy 실행 전 볼륨 값으로 복구", "시스템 볼륨 제어 가능 환경에서만 보장"],
            ["메신저/메일 앱", "Privacy가 숨긴 앱을 다시 표시 또는 활성화 시도", "앱 종료, 권한, Space 상태에 따라 완전 복구가 제한될 수 있음"],
            ["업무 모드", "Privacy 실행 전 업무 모드/Focus 상태로 복구 시도", "macOS Focus 제어 방식과 권한에 따라 제한"],
            ["미디어", "Privacy가 Pause를 보낸 경우 이전 재생 상태 복구는 기본 범위에서 제외", "어떤 콘텐츠가 재생 중이었는지 완전 식별하지 않음"],
            ["창 표시 상태", "Privacy 실행 전 활성 앱/창을 다시 활성화하고 가능한 창 상태 복구", "숨김 해제와 활성화 중심, 모든 창 위치/Space 복원은 보장하지 않음"],
        ],
        [1900, 3300, 4160],
    )
    add_heading(doc, "상태 저장 원칙", 3)
    add_bullet(doc, "Flick Privacy 실행 직전에 복구용 스냅샷을 저장한다.")
    add_bullet(doc, "스냅샷에는 밝기, 볼륨, 업무 모드 상태, Privacy가 숨긴 앱 목록, Privacy 실행 전 활성 앱/창 정보를 포함한다.")
    add_bullet(doc, "Flick Privacy Cancel은 가장 최근 Flick Privacy 스냅샷이 있을 때만 의미 있게 동작한다.")
    add_bullet(doc, "스냅샷이 없거나 일부 항목을 복구할 수 없으면 가능한 항목만 복구하고 사용자에게 요약한다.")
    add_heading(doc, "수용 기준", 3)
    add_number(doc, "Close에 Flick Privacy가 할당되면 Open에는 기본적으로 Flick Privacy Cancel이 자동 할당된다.")
    add_number(doc, "사용자는 Open의 기본 할당을 다른 기능으로 수정할 수 있다.")
    add_number(doc, "Open 제스처로 Flick Privacy Cancel이 실행되면 가장 최근 Privacy 실행 전 밝기와 볼륨으로 복구를 시도한다.")
    add_number(doc, "Privacy가 숨긴 앱은 가능한 범위에서 다시 표시 또는 활성화된다.")
    add_number(doc, "업무 모드는 가능한 범위에서 Privacy 실행 전 상태로 복구된다.")
    add_number(doc, "복구할 스냅샷이 없거나 일부 복구에 실패해도 앱이 크래시하지 않고 결과를 요약한다.")

    add_heading(doc, "7. 우선순위 및 개발 순서", 1)
    add_para(doc, "MVP 확장 개발은 Flick Privacy와 Flick Privacy Cancel의 한 쌍에 집중한다. Flick Privacy는 노트북 화면과 소리가 주변에 노출되는 실제 사용 맥락에서만 장점이 뚜렷하고, Flick Privacy Cancel은 사용자가 보호 모드에서 원래 작업 상태로 빠르게 돌아오게 만드는 복구 동작이다.")
    add_matrix_table(
        doc,
        ["순서", "기능", "판단 근거", "권장 액션"],
        [
            ["1", "Flick Arrange", "P0 완료 기능", "현재 동작 유지 및 신규 명칭 규칙 반영 여부만 확인"],
            ["2", "Flick Privacy", "노트북 환경에서만 장점이 뚜렷하고 데모 가치가 높음", "P1 신규 확정 기능으로 구현"],
            ["3", "Flick Privacy Cancel", "Privacy 적용 이후 원복 경험을 완성하는 필수 보조 기능", "Open 기본 할당으로 구현"],
            ["4", "Flick Focus", "Privacy의 현재 창 유지/나머지 숨김 세부 기능으로 재사용", "독립 기능보다 내부 로직으로 우선 활용"],
        ],
        [900, 1800, 4560, 2100],
    )

    add_heading(doc, "8. MVP 범위", 1)
    add_heading(doc, "포함", 2)
    add_bullet(doc, "Close 제스처를 이용한 Flick Arrange 및 Flick Privacy 실행")
    add_bullet(doc, "Close에 Flick Privacy 할당 시 Open에 Flick Privacy Cancel 기본 자동 할당")
    add_bullet(doc, "Open의 기본 할당을 사용자가 자유롭게 다른 기능으로 수정할 수 있는 인터페이스")
    add_bullet(doc, "기능명 Flick 접두어 유지")
    add_bullet(doc, "제스처명 Close/Open 규칙 적용")
    add_bullet(doc, "Flick Privacy 세부 기능별 on/off 설정")
    add_bullet(doc, "Flick Privacy 실행 전 상태 스냅샷 저장 및 Cancel 복구")
    add_bullet(doc, "밝기 목표값 0~100% 슬라이더와 기본 30% 설정")
    add_bullet(doc, "메신저/메일 기본 숨김 앱 목록 및 사용자 추가/삭제 목록")
    add_bullet(doc, "Flick Focus 로직 재사용을 통한 현재 창 외 나머지 숨김")
    add_heading(doc, "제외", 2)
    add_bullet(doc, "Open 제스처 기반 Launcher, Recent Apps, Recall 기능")
    add_bullet(doc, "Flick Focus 및 Flick Hide의 독립 신규 기능화")
    add_bullet(doc, "시스템 UI, Dock, 메뉴바 제어")
    add_bullet(doc, "고급 창 관리 앱 수준의 레이아웃 커스터마이징")
    add_bullet(doc, "재생 중인 미디어 앱/콘텐츠의 완전한 자동 식별")
    add_bullet(doc, "모든 메신저 앱의 의미론적 자동 분류")

    add_heading(doc, "9. UX 및 데모 시나리오", 1)
    add_para(doc, "데모는 사용자가 키보드 단축키를 외우지 않아도 노트북 화면을 닫는 동작만으로 작업 환경이 바뀐다는 점을 보여주는 데 집중한다.")
    add_number(doc, "여러 앱 창이 열린 복잡한 작업 상태를 준비한다.")
    add_number(doc, "카페나 회의실처럼 주변 노출이 있는 노트북 작업 상황을 준비한다.")
    add_number(doc, "메신저, 메일, 브라우저, 작업 앱, 미디어 재생 상태를 함께 열어둔다.")
    add_number(doc, "Flick Privacy 설정에서 6개 세부 기능이 켜져 있는지 확인한다.")
    add_number(doc, "Close 제스처로 Flick Privacy를 실행한다.")
    add_number(doc, "화면 밝기, 볼륨, 민감 앱 숨김, 업무 모드, 미디어 일시정지, 활성 창 유지 상태를 확인한다.")
    add_number(doc, "Open 제스처로 Flick Privacy Cancel을 실행한다.")
    add_number(doc, "밝기, 볼륨, 업무 모드, 숨김 앱, 활성 앱/창 상태가 Privacy 실행 이전 기준으로 가능한 범위에서 복구되는지 확인한다.")

    add_heading(doc, "10. 검수 체크리스트", 1)
    add_matrix_table(
        doc,
        ["검수 항목", "기준", "결과"],
        [
            ["제스처 명칭", "입력 제스처 문구에 Flick 단어를 사용하지 않는다.", "확인 필요"],
            ["기능 명칭", "앱 기능명은 Flick Arrange, Flick Privacy로 표기한다.", "확인 필요"],
            ["Open 기본 할당", "Close에 Flick Privacy가 할당되면 Open에 Flick Privacy Cancel이 기본 할당된다.", "확인 필요"],
            ["Open 수정 가능", "사용자가 Open에 할당된 Flick Privacy Cancel을 다른 기능으로 변경할 수 있다.", "확인 필요"],
            ["세부 설정", "Flick Privacy의 6개 세부 기능을 각각 on/off 할 수 있다.", "확인 필요"],
            ["상태 스냅샷", "Flick Privacy 실행 전 복구용 상태를 저장한다.", "확인 필요"],
            ["밝기", "기본값은 전체 밝기 기준 30%이며 슬라이더로 조절 가능하다.", "확인 필요"],
            ["볼륨", "볼륨 기능이 켜져 있으면 시스템 볼륨이 0%가 된다.", "확인 필요"],
            ["민감 앱 숨김", "기본 메신저/메일 앱과 사용자 지정 앱을 숨길 수 있다.", "확인 필요"],
            ["업무 모드", "업무 모드 전환을 시도하고 제한이 있으면 사용자에게 안내한다.", "확인 필요"],
            ["미디어", "현재 시스템 미디어 재생에 Pause 명령을 보낸다.", "확인 필요"],
            ["Flick Focus 재사용", "활성 창 1개만 유지하고 나머지 일반 앱 창을 숨긴다.", "확인 필요"],
            ["Cancel 복구", "Open 제스처로 Privacy 이전 밝기, 볼륨, 업무 모드, 숨김 앱, 활성 창 상태 복구를 시도한다.", "확인 필요"],
            ["시스템 UI", "메뉴바, Dock, 시스템 UI는 제어 대상에서 제외한다.", "확인 필요"],
            ["반복 실행", "동일 제스처 반복 실행 시 오류나 앱 상태 깨짐이 없어야 한다.", "확인 필요"],
        ],
        [2400, 5160, 1800],
    )

    add_heading(doc, "11. 향후 확장 아이디어", 1)
    add_para(doc, "Open 제스처는 현재 Flick Privacy Cancel의 기본 실행 트리거로 사용한다. Launcher, Recent Apps, Recall 같은 Open 기반 확장 기능과 Flick Focus/Flick Hide의 독립 기능화는 현재 개발 범위에서 제외하며, Flick Privacy / Cancel 검증 이후 별도 기획으로 전환한다.")
    add_bullet(doc, "Launcher: 자주 쓰는 앱 또는 액션 실행")
    add_bullet(doc, "Recent Apps: 최근 사용 앱 전환")
    add_bullet(doc, "Recall: 이전 작업 상태 복원 또는 최근 작업 호출")

    add_heading(doc, "12. 결론", 1)
    add_para(
        doc,
        "Flick Series의 MVP 확장 방향은 창 관리 기능 자체보다 ‘노트북 힌지를 입력 장치로 사용할 수 있는가’라는 제품 가설 검증에 초점을 둔다. Flick Privacy는 노트북 화면과 소리가 주변에 노출되는 순간을 겨냥하므로, 외부 모니터 고정 환경보다 노트북 환경에서 장점이 선명하다. 초기 구현은 Close 제스처 한 번으로 사용자가 선택한 프라이버시 보호 동작들을 안정적으로 실행하고, Open 제스처 한 번으로 Privacy 이전 환경을 가능한 한 정확히 복구하는 데 집중한다.",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
